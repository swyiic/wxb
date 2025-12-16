# -*- coding: utf-8 -*-
import sys
import json
from docx import Document
from docx.shared import Inches
import base64
import io
import os
from bs4 import BeautifulSoup
import time

# 确保中文不乱码
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# ==================== 你原来的所有函数保持不变 ====================
def add_padding_if_needed(base64_string):
    return base64_string + "=" * ((4 - len(base64_string) % 4) % 4)

def parse_html_content(html_content):
    if not html_content:
        return []
    html_content = html_content.replace('\\"', '"')
    if not html_content.strip().startswith('<'):
        html_content = f'<div>{html_content}</div>'
    soup = BeautifulSoup(html_content, 'html.parser')
    elements = []
    for child in soup.descendants:
        if child.name is None and child.string and child.string.strip():
            elements.append({'type': 'text', 'content': child.string.strip()})
        elif child.name == 'img' and 'src' in child.attrs:
            src = child.attrs['src']
            if src.startswith('data:image'):
                base64_data = src.split('base64,')[-1]
                base64_data = add_padding_if_needed(base64_data)
                try:
                    img_bytes = base64.b64decode(base64_data)
                    elements.append({'type': 'image', 'data': io.BytesIO(img_bytes)})
                except:
                    elements.append({'type': 'text', 'content': '[图片解析失败]'})
        elif child.name in ['br', 'p', 'div']:
            elements.append({'type': 'break'})
    return elements

def replace_placeholders_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text or f"{{{placeholder}}}" in run.text or f"[{placeholder}]" in run.text:
                run.text = run.text.replace(placeholder, "").replace(f"{{{placeholder}}}", "").replace(f"[{placeholder}]", "")
                parsed_elements = parse_html_content(value)
                current_run = run
                last_type = None
                for element in parsed_elements:
                    if element['type'] == 'text':
                        if last_type in ['image', 'break']:
                            current_run.add_break()
                        current_run.text += element['content']
                        last_type = 'text'
                    elif element['type'] == 'image':
                        current_run = paragraph.add_run()
                        current_run.add_picture(element['data'], width=Inches(4))
                        current_run = paragraph.add_run()
                        last_type = 'image'
                    elif element['type'] == 'break':
                        if last_type != 'break' and last_type is not None:
                            current_run.add_break()
                            last_type = 'break'
                break

def replace_text_and_images(doc, replacements):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)
# ============================================================

def generate_report(json_path):
    # 1. 读取 Rust 传来的临时 json
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    replacements = {
        "CN": data.get("company_name", ""),
        "CU": data.get("web_url", ""),
        "WT": data.get("system_name", ""),
        "VN": data.get("vuln_name", ""),
        "CD": data.get("find_date", ""),
        "VU": data.get("vuln_url", ""),
        "VD": data.get("vuln_description", ""),
        "FS": data.get("fix_suggestion", ""),
        "affectsType": data.get("affects_type",""),
        "AP": data.get("assist_image", ""),
        "TR": data.get("test_process", ""),
        "AD": data.get("appendix", "")
    }

    # 脚本所在目录（打包后也是 exe 同目录）
    if getattr(sys, "frozen", False):
        script_dir = os.path.dirname(sys.executable)
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.join(script_dir, "demo.docx")
    reports_dir = os.path.join(script_dir, "reports")
    pending_dir = os.path.join(script_dir, "pending_reports")
    os.makedirs(reports_dir, exist_ok=True)
    os.makedirs(pending_dir, exist_ok=True)

    doc = Document(template_path)
    replace_text_and_images(doc, replacements)

    # 生成报告编号
    existing = [f for f in os.listdir(reports_dir) if f.endswith("-网络安全检测报告.docx")]
    numbers = [int(f.split('-')[0]) for f in existing if f.split('-')[0].isdigit()]
    next_num = max(numbers, default=0) + 1
    docx_path = os.path.join(reports_dir, f"{next_num}-{replacements['CN']}-网络安全检测报告.docx")
    doc.save(docx_path)

    # ==================== 关键：生成你完全认识的 middle_json ====================
    middle_json = {
        "companyName": replacements["CN"],
        "companyWebName": replacements["WT"],
        "province": "31",
        "city": "3101",
        "district": "",
        "address": "",
        "alarmTargetIp": "",
        "alarmPort": "",
        "alarmTargetUrl": replacements["CU"],
        "alarmSort": 2,
        "riskType": 0,            # Rust 会自动映射成 121 等
        "informationName": replacements["VN"],
        "foundTime": replacements["CD"],
        "alarmDesc": replacements["VD"],
        "affectsType": replacements["affectsType"],
        "solution": replacements["FS"],
        "fileIds": [],
        "taskId": "",
        "taskListId": "",

        # 下面两个字段 Rust 必须用到
        "localReportPath": os.path.abspath(docx_path),   # DOCX 绝对路径
        "_pending_file": ""       # 马上填自己的路径
    }

    # 保存到 pending_reports 目录
    timestamp = int(time.time() * 1000)
    pending_path = os.path.join(pending_dir, f"pending_{next_num}_{timestamp}.json")
    middle_json["_pending_file"] = os.path.abspath(pending_path)

    with open(pending_path, 'w', encoding='utf-8') as f:
        json.dump(middle_json, f, ensure_ascii=False, indent=2)

    # ==================== 最后只打印这个路径，和你原来完全一样 ====================
    print(os.path.abspath(pending_path))
    sys.stdout.flush()

if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit(1)
    generate_report(sys.argv[1])